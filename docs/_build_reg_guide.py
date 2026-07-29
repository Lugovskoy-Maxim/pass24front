# -*- coding: utf-8 -*-
"""Generate tenant registration instruction DOCX."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name("ИНСТРУКЦИЯ_РЕГИСТРАЦИЯ_АРЕНДАТОРА.docx")


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Arial"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:eastAsia"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def add_runs(paragraph, text: str, bold=False, italic=False, size=None, color=None):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, italic=italic, color=color)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, bold=bold, italic=italic, color=color)


def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.15

    for level, size, before, after in ((1, 18, 6, 10), (2, 14, 14, 6), (3, 12, 10, 4)):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Arial"
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        hs.paragraph_format.space_before = Pt(before)
        hs.paragraph_format.space_after = Pt(after)

    def p(text="", bold=False, italic=False, size=None, space_after=6, align=None, color=None):
        para = doc.add_paragraph()
        if align is not None:
            para.alignment = align
        para.paragraph_format.space_after = Pt(space_after)
        if text:
            add_runs(para, text, bold=bold, italic=italic, size=size, color=color)
        return para

    def bullet(text):
        para = doc.add_paragraph(style="List Bullet")
        # clear default empty run content carefully
        if para.runs:
            for r in list(para.runs):
                r.text = ""
        add_runs(para, text, size=11)
        para.paragraph_format.space_after = Pt(3)

    def numbered(text):
        para = doc.add_paragraph(style="List Number")
        if para.runs:
            for r in list(para.runs):
                r.text = ""
        add_runs(para, text, size=11)
        para.paragraph_format.space_after = Pt(3)

    def h2(t):
        doc.add_heading(t, level=2)

    def h3(t):
        doc.add_heading(t, level=3)

    def table(headers, rows, col_widths=None):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(h)
            set_run_font(run, size=10, bold=True)
            set_cell_shading(cell, "E8EEF7")
            set_cell_margins(cell)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.rows[ri + 1].cells[ci]
                cell.text = ""
                para = cell.paragraphs[0]
                add_runs(para, str(val), size=10)
                set_cell_margins(cell)
        if col_widths:
            for row in t.rows:
                for i, w in enumerate(col_widths):
                    row.cells[i].width = Cm(w)
        doc.add_paragraph()

    # Title
    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t1.add_run("Инструкция")
    set_run_font(r, size=12, bold=True, color=RGBColor(0x4A, 0x55, 0x68))

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t2.add_run("Регистрация арендатора в Pass")
    set_run_font(r, size=20, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E))

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(t3, "Система электронных пропусков Pass · https://pass.mstyle.ru", size=11)
    t3.paragraph_format.space_after = Pt(4)

    t4 = doc.add_paragraph()
    t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(
        t4,
        "Для владельцев компаний-арендаторов (не для сотрудников БЦ и не для охраны)",
        italic=True,
        size=10,
        color=RGBColor(0x4A, 0x55, 0x68),
    )
    t4.paragraph_format.space_after = Pt(12)

    h2("Кому подходит эта инструкция")
    table(
        ["Роль", "Как появляется в системе"],
        [
            ["Арендатор (владелец компании)", "Самостоятельно регистрируется на сайте"],
            ["Сотрудник компании", "Приглашает владелец из раздела «Профиль» (ссылка на email)"],
            ["Ресепшн / админ БЦ", "Создаёт администратор"],
        ],
        col_widths=[7.5, 9.0],
    )
    p(
        "Если вы **сотрудник** арендатора — попросите владельца компании отправить приглашение "
        "из профиля. Эта инструкция для **самостоятельной регистрации владельца**."
    )

    h2("Что нужно подготовить")
    numbered("**Корпоративная почта** в зоне **.ru**, **.рф** или **.su** (пример: ivanov@company.ru)")
    numbered("**Название компании** — как в договоре / на табличке (например, ООО «Ромашка»)")
    numbered("**ФИО** ответственного лица")
    numbered("**Пароль** — не короче 6 символов")
    numbered("По желанию — **мобильный телефон** в формате +7… (для входа и SMS-кода, если SMS включено)")

    h3("Какая почта не подойдёт")
    p("Обычно **не принимаются** публичные сервисы:")
    bullet("Gmail, Outlook / Hotmail, Yahoo, iCloud, ProtonMail и аналоги")
    bullet("Адреса **не** в зоне .ru / .рф / .su")
    p("Используйте рабочую корпоративную почту компании.")

    h2("Шаг 1. Откройте страницу входа")
    numbered("Перейдите на сайт Pass: **https://pass.mstyle.ru**")
    numbered("Откроется страница входа (или перейдите по ссылке «Вход» /login)")
    numbered("Выберите режим **«Регистрация»** (не «Вход»)")

    h2("Шаг 2. Заполните анкету")
    p("Укажите:")
    table(
        ["Поле", "Обязательно", "Комментарий"],
        [
            ["Фамилия", "да", ""],
            ["Имя", "да", ""],
            ["Отчество", "нет", "по желанию"],
            ["Компания", "да", "полное или краткое юр. название"],
            ["Email", "да*", "корпоративный, см. требования выше"],
            ["Телефон", "нет / да*", "формат +7 9XX XXX-XX-XX"],
            ["Пароль", "да", "минимум 6 символов"],
            ["Повтор пароля", "да", "должен совпасть с паролем"],
        ],
        col_widths=[4.0, 3.5, 9.0],
    )
    p("* Канал подтверждения:")
    bullet("**По email** — код придёт на указанную почту (основной способ)")
    bullet("**По SMS** — если функция включена администратором; код на телефон (не чаще 1 раза в 5 минут)")
    p("Нажмите кнопку отправки кода / продолжения регистрации.")

    h2("Шаг 3. Подтвердите код")
    numbered("Откройте письмо (или SMS) с **6-значным кодом**")
    numbered("Код действует **15 минут**")
    numbered("Введите код на сайте")
    numbered("Если письма нет — проверьте папку **«Спам»** / «Промоакции»")
    numbered("Если код истёк — запросите код **заново** (для SMS — не чаще чем раз в 5 минут)")
    p("После успешного ввода кода **учётная запись создаётся**.")

    h2("Шаг 4. Ожидание активации администратором")
    p("Сразу после регистрации доступ **ещё не полный**:")
    bullet("Заявка попадает администратору на проверку")
    bullet("Пока заявка **не одобрена**, заказ пропусков **недоступен**")
    bullet("На сайте может отображаться сообщение: «Заявка на регистрацию ожидает подтверждения»")

    h3("Что делает администратор")
    numbered("Проверяет данные компании и контакты")
    numbered("**Одобряет** регистрацию")
    numbered("**Привязывает офисы** компании к вашему аккаунту")
    p("Без закреплённых офисов заказ пропусков также будет недоступен.")

    h2("Шаг 5. Первый вход после одобрения")
    numbered("Откройте **https://pass.mstyle.ru/login**")
    numbered("Войдите по **email**, **телефону** (+7…) или логину (если выдан)")
    numbered("Введите **пароль**, заданный при регистрации")
    p("После входа станут доступны разделы арендатора:")
    bullet("заказ пропусков")
    bullet("список своих пропусков")
    bullet("профиль и (для владельца) приглашение сотрудников")

    h2("Что делать дальше (кратко)")

    h3("Заказать пропуск")
    numbered("**Пропуска** → **Заказать пропуск**")
    numbered("Тип пропуска, данные гостя, дата (ближайшие рабочие дни), офис")
    numbered("При необходимости — авто (госномер), комментарий, отправка ссылки гостю на email")
    numbered("Отправить заявку → открывается **электронный пропуск** со ссылкой и QR")

    h3("Добавить сотрудника компании")
    numbered("**Профиль** → блок **«Сотрудники компании»**")
    numbered("ФИО, email сотрудника → **Отправить приглашение**")
    numbered("Сотрудник переходит по ссылке из письма (действует **72 часа**) и задаёт пароль")
    numbered("Лимит: **не более 3 сотрудников** на компанию (кроме владельца)")

    h3("Если забыли пароль")
    numbered("На странице входа — **«Забыли пароль?»**")
    numbered("Укажите **email** аккаунта")
    numbered("Введите код из письма и новый пароль")
    p(
        "Если email не подходит / не помните — обратитесь к администратору БЦ "
        "(контакты на странице входа или в «Помощь»)."
    )

    h2("Частые вопросы")

    h3("«Не приходит код»")
    bullet("Проверьте **Спам**")
    bullet("Убедитесь, что email **корпоративный** и в зоне .ru / .рф / .su")
    bullet("Дождитесь **15 минут** и запросите код снова")
    bullet("Для SMS: не чаще **1 раза в 5 минут**; если SMS отключено — регистрируйтесь **по email**")

    h3("«Регистрация отклоняет email»")
    bullet("Нельзя Gmail, Outlook, iCloud и другие публичные сервисы")
    bullet("Нужна почта вида имя@ваша-компания.ru")

    h3("«Вошёл, но не могу заказать пропуск»")
    numbered("Заявка ещё **не одобрена** администратором — дождитесь активации")
    numbered("Одобрили, но **не привязали офис** — напишите администратору БЦ")
    numbered("Вы **сотрудник** без прав / отключён владельцем — уточните у владельца компании")

    h3("«Нужен доступ второму человеку из компании»")
    p(
        "Не создавайте вторую регистрацию компании. "
        "Владелец добавляет сотрудника в **Профиль → Сотрудники**."
    )

    h3("«Мы уже зарегистрированы на сайте M-Style»")
    p(
        "Регистрация в **Pass** — отдельная. Даже если вы резидент на сайте управляющей компании, "
        "в Pass нужна своя учётная запись (или доступ, который выдаст администратор)."
    )

    h2("Контакты при проблемах")
    bullet("Контакты на странице **входа** / в блоке **«Помощь»** в приложении")
    bullet("Ресепшн / администратор бизнес-центра")
    bullet("Управляющая компания (если так принято в вашем БЦ)")
    p("Укажите при обращении:")
    bullet("название **компании**")
    bullet("**email**, с которого регистрировались")
    bullet("**телефон** (если указывали)")
    bullet("что именно не получается (скрин ошибки)")

    h2("Краткая схема")
    for line in (
        "1. pass.mstyle.ru → Регистрация",
        "2. ФИО + компания + корпоративный email + пароль",
        "3. Код из письма / SMS (15 мин)",
        "4. Ожидание одобрения администратором + привязка офисов",
        "5. Вход → заказ пропусков",
    ):
        p(line, space_after=2)

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(16)
    add_runs(
        foot,
        "Документ описывает регистрацию арендатора (владельца компании) в системе Pass. "
        "Сотрудники компании подключаются только по приглашению владельца.",
        italic=True,
        size=9,
        color=RGBColor(0x66, 0x66, 0x66),
    )

    # Page number footer
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Pass · регистрация арендатора · стр. ")
    set_run_font(run, size=9, color=RGBColor(0x88, 0x88, 0x88))

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run2 = fp.add_run()
    run2._r.append(fld_begin)
    run2._r.append(instr)
    run2._r.append(fld_end)
    set_run_font(run2, size=9, color=RGBColor(0x88, 0x88, 0x88))

    doc.save(str(OUT))
    print(f"OK {OUT}")
    print(f"size {OUT.stat().st_size}")


if __name__ == "__main__":
    main()
